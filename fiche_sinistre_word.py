# -*- coding: utf-8 -*-
"""
Génération de la fiche de sinistre au format Word (.docx) par remplacement
de balises (ex: {{numero_fiche}}, {{date_sinistre}}, etc.).

Avantages majeurs du format Word par rapport à la superposition PDF :
- ZÉRO calibrage, zéro coordonnée (x, y) à régler : Word gère automatiquement
  les retours à la ligne, les tableaux, les polices et les marges.
- L'utilisateur peut personnaliser son modèle (FICHE_DE_SINISTRE_MODELE.docx)
  directement dans Microsoft Word en y insérant des balises {{...}}.
- Le fichier généré est 100% modifiable par l'utilisateur dans Word avant impression.

Balises supportées dans le document Word :
  {{numero_fiche}}         (ex: 009/2026)
  {{date_sinistre}}        (ex: 08/08/2026)
  {{lieu_accident}}        (ex: Alger)
  {{immatriculation}}      (ou {{matricule_vehicule}})
  {{chauffeur}}
  {{pv_recu}}              (ou {{pv_autorites}})
  {{autorite_pv}}          (ou {{autorite}})
  {{adresse_autorite}}
  {{avec_sans_tiers}}      (ou {{tiers}})
  {{documents_recuperes}}
  {{degats_cause}}         (ou {{degats}})
  {{circonstance_accident}} (ou {{circonstances}})
"""
import os
import logging
from pathlib import Path
import database as db

logger = logging.getLogger(__name__)

TEMPLATE_WORD_FILENAME = "FICHE_DE_SINISTRE_MODELE.docx"


def get_word_template_path():
    """Localise le modèle Word (.docx) dans le dossier de données de l'application
    ou dans le dossier du script. Retourne le chemin ou None."""
    base_dir = Path(__file__).resolve().parent
    candidates = [
        os.path.join(db.get_app_dir(), TEMPLATE_WORD_FILENAME),
        str(base_dir / TEMPLATE_WORD_FILENAME),
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    return None


def word_template_is_available():
    """True si un modèle Word (FICHE_DE_SINISTRE_MODELE.docx) est présent."""
    return get_word_template_path() is not None


def convert_doc_to_docx(doc_path, dest_dir=None):
    """Convertit automatiquement un fichier Word 97-2003 (.doc) en format moderne
    (.docx) via Microsoft Word (win32com) ou LibreOffice/OpenOffice si disponible.
    Retourne le chemin du fichier .docx généré ou None en cas d'échec."""
    if dest_dir is None:
        dest_dir = db.get_app_dir()
    base_name = os.path.splitext(os.path.basename(doc_path))[0]
    out_docx = os.path.join(dest_dir, base_name + ".docx")
    if os.path.abspath(doc_path) == os.path.abspath(out_docx):
        return out_docx

    # 1. Tentative avec Microsoft Word sous Windows (win32com)
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(doc_path))
        doc.SaveAs2(os.path.abspath(out_docx), FileFormat=16)  # 16 = wdFormatXMLDocument (.docx)
        doc.Close()
        word.Quit()
        if os.path.exists(out_docx):
            return out_docx
    except Exception as e:
        logger.debug("win32com conversion doc -> docx failed: %s", e)

    # 2. Tentative avec LibreOffice / OpenOffice en ligne de commande
    try:
        import subprocess, shutil
        executables = ["soffice", "libreoffice",
                       r"C:\Program Files\LibreOffice\program\soffice.exe",
                       r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"]
        for exe in executables:
            if shutil.which(exe) or os.path.exists(exe):
                subprocess.run([exe, "--headless", "--convert-to", "docx", doc_path,
                                "--outdir", dest_dir], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                if os.path.exists(out_docx):
                    return out_docx
    except Exception as e:
        logger.debug("LibreOffice conversion doc -> docx failed: %s", e)

    return None


def import_user_word_template(src_path):
    """Importe le fichier Word choisi par l'utilisateur (.doc ou .docx) dans le
    dossier de données de l'application comme FICHE_DE_SINISTRE_MODELE.docx.

    Si le fichier source est un .doc (Word 97-2003), tente de le convertir en .docx.
    Retourne (succes: bool, chemin_dest: str, message: str).
    """
    import shutil
    dest_dir = db.get_app_dir()
    dest_path = os.path.join(dest_dir, TEMPLATE_WORD_FILENAME)

    if src_path.lower().endswith(".doc") and not src_path.lower().endswith(".docx"):
        converted = convert_doc_to_docx(src_path, dest_dir=dest_dir)
        if converted and os.path.exists(converted):
            shutil.copy2(converted, dest_path)
            return True, dest_path, ("Votre fichier Word ancien format (.doc) a été converti et importé "
                                     "avec succès au format moderne (.docx) !")
        else:
            return False, "", ("Votre fichier est au format ancien Word 97-2003 (.doc) et la conversion "
                               "automatique a échoué (Microsoft Word / LibreOffice non détectés).\n\n"
                               "Veuillez ouvrir votre fichier une fois dans Microsoft Word, faire "
                               "« Fichier > Enregistrer sous » en choisissant « Document Word (*.docx) », "
                               "puis recharger ce fichier .docx.")
    else:
        shutil.copy2(src_path, dest_path)
        return True, dest_path, ("Le modèle Word a été importé avec succès :\n" + dest_path)


def _get_replacement_dict(data):
    """Construit un dictionnaire associant chaque balise {{...}} à sa valeur."""
    import fiche_sinistre_template as tpl
    def val(key):
        v = data.get(key)
        return "" if v in (None, "") else str(v).strip()

    num_plain = tpl._fiche_number_plain(data) if hasattr(tpl, "_fiche_number_plain") else str(data.get("fiche_number", ""))
    num_full = str(data.get("fiche_number", num_plain))

    mapping = {
        "{{numero_fiche}}": num_plain,
        "{{numero_fiche_complet}}": num_full,
        "{{code_cam}}": val("code_cam"),
        "{{cam}}": val("code_cam"),
        "{{date_sinistre}}": val("date_sinistre"),
        "{{lieu_accident}}": val("lieu_accident"),
        "{{immatriculation}}": val("immatriculation"),
        "{{matricule_vehicule}}": val("immatriculation"),
        "{{chauffeur}}": val("chauffeur"),
        "{{pv_recu}}": val("pv_recu"),
        "{{pv_autorites}}": val("pv_recu"),
        "{{autorite_pv}}": val("autorite_pv"),
        "{{autorite}}": val("autorite_pv"),
        "{{adresse_autorite}}": val("adresse_autorite"),
        "{{avec_sans_tiers}}": val("avec_sans_tiers"),
        "{{tiers}}": val("avec_sans_tiers"),
        "{{documents_recuperes}}": val("documents_recuperes"),
        "{{degats_cause}}": val("degats_cause"),
        "{{degats}}": val("degats_cause"),
        "{{circonstance_accident}}": val("circonstance_accident"),
        "{{circonstances}}": val("circonstance_accident"),
    }
    return mapping


def _replace_tags_in_paragraph(p, mapping):
    """Remplace les balises {{...}} dans un paragraphe Word (docx.Paragraph).
    Gère le cas où Word découpe une balise sur plusieurs runs."""
    full_text = p.text
    if not full_text or "{{" not in full_text:
        return False

    new_text = full_text
    for tag, val in mapping.items():
        if tag in new_text:
            new_text = new_text.replace(tag, val)

    if new_text != full_text:
        # On affecte le texte remplacé au premier run pour conserver sa mise en forme
        if len(p.runs) > 0:
            p.runs[0].text = new_text
            for run in p.runs[1:]:
                run.text = ""
        else:
            p.add_run(new_text)
        return True
    return False


def _replace_tags_in_table(table, mapping):
    """Parcourt les cellules et sous-tableaux d'un tableau Word pour remplacer les balises."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _replace_tags_in_paragraph(p, mapping)
            for sub_table in cell.tables:
                _replace_tags_in_table(sub_table, mapping)


import unicodedata
import re

KEYWORD_PATTERNS = [
    ("code_cam", ["code cam", "n code cam", "numero code cam", "code c a m", "cam"]),
    ("date_sinistre", ["date de sinistre", "date du sinistre", "date accident", "date d accident", "date sinistre", "survenu le"]),
    ("lieu_accident", ["lieu d accident", "lieu de l accident", "lieu accident", "lieu du sinistre", "lieu sinistre", "endroit"]),
    ("immatriculation", ["matricule", "immatriculation", "vehicule", "n d immatriculation", "numero d immatriculation"]),
    ("chauffeur", ["chauffeur", "conducteur", "nom du chauffeur", "nom et prenom du chauffeur", "conduit par"]),
    ("pv_recu", ["y a t il un pv", "pv des autorites", "p v des autorites", "pv autorite", "proces verbal", "pv recu", "p v recu"]),
    ("adresse_autorite", ["adresse de l autorite", "adresse autorite", "brigade commissariat", "adresse de la brigade", "lieu de l autorite"]),
    ("autorite_pv", ["quelle autorite", "gendarmerie police", "autorite du pv", "autorite pv", "police ou gendarmerie", "etabli par", "autorite"]),
    ("avec_sans_tiers", ["avec ou sans tiers", "avec sans tiers", "tiers implique", "tiers", "impliquant un tiers"]),
    ("documents_recuperes", ["copies des documents", "documents recuperes", "copies recuperees", "pieces recuperees", "si oui les copies", "documents"]),
    ("degats_cause", ["degats causes", "degats materiels", "nature des degats", "degats au vehicule", "degats", "dommages"]),
    ("circonstance_accident", ["circonstances de l accident", "circonstance de l accident", "circonstances d accident", "circonstance", "description de l accident", "deroulement", "circonstances"]),
]


def _normalize_text(s):
    if not s:
        return ""
    s = str(s).lower()
    s = ''.join(c for c in unicodedata.normalize('NFD', s)
                if unicodedata.category(c) != 'Mn')
    s = re.sub(r'[^a-z0-9\s]', ' ', s)
    s = re.sub(r'\s+', ' ', s).strip()
    return s


def _match_field_key(norm_text):
    """Retourne la clé de champ correspondant au texte normalisé du libellé."""
    if not norm_text or len(norm_text) < 3:
        return None
    for field_key, keywords in KEYWORD_PATTERNS:
        for kw in keywords:
            if kw in norm_text:
                return field_key
    return None


def _safe_replace_or_set_text(p, new_text, strip_bold_underline=True):
    """Remplace le texte d'un paragraphe docx.Paragraph en conservant 100% de la
    taille et police d'origine, mais sans mettre en gras ni souligner
    (le texte inséré reste dans l'écriture normale du fichier d'origine)."""
    if len(p.runs) > 0:
        r = p.runs[0]
        r.text = str(new_text)
        if strip_bold_underline:
            r.bold = False
            r.underline = False
        for extra in p.runs[1:]:
            extra.text = ""
    else:
        r = p.add_run(str(new_text))
        if strip_bold_underline:
            r.bold = False
            r.underline = False


def _fill_paragraph_if_match(p, filled_fields, get_val_fn):
    """Analyse un paragraphe et le remplit si c'est un libellé suivi de pointillés/deux-points,
    en conservant 100% de la mise en forme d'origine (police, taille, gras...)."""
    text = p.text.strip()
    if not text or len(text) < 4:
        return
    norm = _normalize_text(text)
    field_key = _match_field_key(norm)
    if field_key and field_key not in filled_fields:
        val = get_val_fn(field_key)
        if not val:
            return
        if ":" in p.text:
            label_part = p.text.split(":", 1)[0].strip()
            _safe_replace_or_set_text(p, f"{label_part} : {val}")
            filled_fields.add(field_key)
        elif ".." in p.text or "__" in p.text:
            label_part = re.split(r"(\.\.+|__+)", p.text)[0].strip()
            _safe_replace_or_set_text(p, f"{label_part} : {val}")
            filled_fields.add(field_key)


def auto_scan_and_fill_word_doc(doc, data):
    """Système intelligent de scan et remplissage automatique d'un document Word
    (.docx) de fiche de sinistre NE CONTENANT PAS de balises {{...}}.

    1. Parcourt tous les tableaux (doc.tables) :
       - Identifie intelligemment quelle ligne correspond à quel renseignement
         en analysant le texte de la colonne de gauche (ex: 'Date de sinistre :',
         'Chauffeur', 'Immatriculation', 'PV', 'Dégâts', etc.).
       - Insère automatiquement la valeur du sinistre dans la colonne de droite (ou
         à la suite du libellé dans la cellule) en conservant 100% de la mise en page
         et des tailles de police d'origine du fichier.
    2. Parcourt tous les paragraphes du document (doc.paragraphs) :
       - Détecte le titre de la fiche ('FICHE DE SINISTRE') et ajoute le n° de fiche.
       - Détecte les libellés sous forme 'Libellé : .......' et insère la valeur après
         le deux-points ':'.
    """
    filled_fields = set()
    mapping = _get_replacement_dict(data)

    def get_val_for_key(field_key):
        if field_key == "numero_fiche":
            return mapping.get("{{numero_fiche}}", "")
        tag = f"{{{{{field_key}}}}}"
        return mapping.get(tag, "")

    # 1. Scan des Tableaux (le cas le plus courant pour un formulaire officiel)
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                left_cell = row.cells[0]
                right_cell = row.cells[-1]  # dernière cellule de la ligne
                norm_left = _normalize_text(left_cell.text)
                field_key = _match_field_key(norm_left)
                if field_key and field_key not in filled_fields:
                    val = get_val_for_key(field_key)
                    if val:
                        if len(right_cell.paragraphs) > 0:
                            _safe_replace_or_set_text(right_cell.paragraphs[0], val)
                            for extra_p in right_cell.paragraphs[1:]:
                                _safe_replace_or_set_text(extra_p, "")
                        else:
                            right_cell.text = val
                        filled_fields.add(field_key)
            elif len(row.cells) == 1:
                _fill_paragraph_if_match(row.cells[0].paragraphs[0], filled_fields, get_val_for_key)

    # 2. Scan des Paragraphes du corps du document
    for p in doc.paragraphs:
        _fill_paragraph_if_match(p, filled_fields, get_val_for_key)

    # 3. Traitement spécial du titre (Numéro de fiche) si non rempli
    num_plain = mapping.get("{{numero_fiche}}", "")
    if num_plain:
        for p in doc.paragraphs[:6]:
            norm_p = _normalize_text(p.text)
            if "fiche de sinistre" in norm_p and num_plain not in p.text:
                if ":" in p.text:
                    new_txt = p.text.split(":", 1)[0] + " : " + num_plain
                elif "n°" in p.text.lower() or "n " in norm_p:
                    new_txt = re.sub(r"(?i)(n°|n)\s*([0-9/.\-_]*)", f"n° {num_plain}", p.text)
                else:
                    new_txt = p.text.strip() + f" n° {num_plain}"
                _safe_replace_or_set_text(p, new_txt)
                break


def build_fiche_word(data, output_path):
    """Génère la fiche de sinistre au format Word (.docx) à partir d'un dict de données.

    Si le modèle FICHE_DE_SINISTRE_MODELE.docx n'existe pas encore, un modèle
    par défaut contenant l'en-tête, le tableau de renseignements et les balises
    est automatiquement créé puis utilisé.
    """
    import docx
    template_path = get_word_template_path()
    if not template_path:
        # Création automatique d'un modèle par défaut complet s'il n'existe pas
        template_path = create_default_word_template()

    doc = docx.Document(template_path)
    mapping = _get_replacement_dict(data)

    # 1. Remplacement dans les paragraphes du corps
    for p in doc.paragraphs:
        _replace_tags_in_paragraph(p, mapping)

    # 2. Remplacement dans les tableaux
    for table in doc.tables:
        _replace_tags_in_table(table, mapping)

    # 3. Remplacement dans les en-têtes et pieds de page des sections
    for section in doc.sections:
        for p in section.header.paragraphs:
            _replace_tags_in_paragraph(p, mapping)
        for table in section.header.tables:
            _replace_tags_in_table(table, mapping)
        for p in section.footer.paragraphs:
            _replace_tags_in_paragraph(p, mapping)
        for table in section.footer.tables:
            _replace_tags_in_table(table, mapping)

    # 4. Scan intelligent et remplissage automatique si le modèle ne contient aucune balise {{...}}
    auto_scan_and_fill_word_doc(doc, data)

    doc.save(output_path)
    return output_path


def create_default_word_template(output_path=None):
    """Crée un modèle Word (.docx) par défaut officiel et soigné, prêt à être
    utilisé tel quel ou personnalisé dans Microsoft Word avec des balises {{...}}."""
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if output_path is None:
        output_path = os.path.join(db.get_app_dir(), TEMPLATE_WORD_FILENAME)

    doc = docx.Document()

    # Configuration des marges de page (2 cm partout)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # En-tête Organisme
    import fiche_sinistre as f_old
    letterhead = f_old.get_letterhead()
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, line in enumerate(letterhead):
        run = p_org.add_run(line + ("\n" if i < len(letterhead) - 1 else ""))
        run.font.name = "Arial"
        if i == 0:
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(31, 58, 95)
        else:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()  # espace

    # Titre de la fiche
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("FICHE DE SINISTRE n° {{numero_fiche}}")
    run_title.bold = True
    run_title.font.name = "Arial"
    run_title.font.size = Pt(16)
    run_title.font.color.rgb = RGBColor(31, 58, 95)

    doc.add_paragraph()  # espace

    # Tableau des informations du sinistre
    fields_def = [
        ("Date de sinistre", "{{date_sinistre}}"),
        ("Lieu d'accident", "{{lieu_accident}}"),
        ("Matricule du véhicule", "{{immatriculation}}"),
        ("Chauffeur", "{{chauffeur}}"),
        ("Y a-t-il un PV des autorités", "{{pv_recu}}"),
        ("Quelle autorité (Gendarmerie / Police)", "{{autorite_pv}}"),
        ("Adresse de l'autorité", "{{adresse_autorite}}"),
        ("Avec ou sans tiers", "{{avec_sans_tiers}}"),
        ("Si oui, copies des documents récupérés", "{{documents_recuperes}}"),
        ("Dégâts causés", "{{degats_cause}}"),
        ("Circonstances de l'accident", "{{circonstance_accident}}"),
    ]

    table = doc.add_table(rows=len(fields_def), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Mise en forme du tableau
    for i, (label_text, tag_text) in enumerate(fields_def):
        row = table.rows[i]
        # Hauteur min
        row.height = Pt(28)
        
        cell_left = row.cells[0]
        cell_right = row.cells[1]
        
        # Largeurs de colonnes (environ 40% / 60%)
        cell_left.width = Inches(2.8)
        cell_right.width = Inches(4.0)

        p_l = cell_left.paragraphs[0]
        p_l.paragraph_format.space_after = Pt(2)
        p_l.paragraph_format.space_before = Pt(2)
        run_l = p_l.add_run(label_text + " :")
        run_l.bold = True
        run_l.font.name = "Arial"
        run_l.font.size = Pt(10.5)
        run_l.font.color.rgb = RGBColor(31, 58, 95)

        p_r = cell_right.paragraphs[0]
        p_r.paragraph_format.space_after = Pt(2)
        p_r.paragraph_format.space_before = Pt(2)
        run_r = p_r.add_run(tag_text)
        run_r.font.name = "Arial"
        run_r.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    # Tableau pour les 2 zones de signature côte à côte
    table_sig = doc.add_table(rows=1, cols=2)
    table_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_sig.style = "Table Grid"
    row_sig = table_sig.rows[0]
    row_sig.height = Inches(1.2)

    for idx, label_sig in enumerate(["Signature du chauffeur", "Signature du responsable"]):
        cell = row_sig.cells[idx]
        cell.width = Inches(3.3)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(50)
        run = p.add_run(label_sig)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(31, 58, 95)

    doc.save(output_path)
    return output_path
