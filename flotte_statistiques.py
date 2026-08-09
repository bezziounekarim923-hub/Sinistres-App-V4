# -*- coding: utf-8 -*-
"""
Module d'analyses de flotte et relevés individuels de sinistralité.
Permet d'extraire, d'analyser et de générer un rapport professionnel Word (.docx)
pour un chauffeur (historique complet, sinistralité, responsabilité) ou pour
un véhicule (immatriculation).
"""
import os
import logging
import datetime
import database as db
import fiche_sinistre as fiche
import date_utils

logger = logging.getLogger(__name__)


def _normalize_name(name):
    """Normalise un nom pour comparaison insensible à la casse et aux espaces."""
    return str(name or "").strip().lower()


def get_driver_statistics(chauffeur_name):
    """Calcule les statistiques et récupère l'historique d'un chauffeur donné."""
    norm_target = _normalize_name(chauffeur_name)
    all_records = db.fetch_all()
    records = [r for r in all_records if _normalize_name(r.get("chauffeur")) == norm_target]

    # Tri par date de sinistre (du plus récent au plus ancien)
    def _sort_key(r):
        dt = date_utils.parse_date_input(r.get("date_sinistre"))
        return dt or datetime.date(1970, 1, 1)

    records.sort(key=_sort_key, reverse=True)

    total = len(records)
    regles = sum(1 for r in records if str(r.get("statut_reglement") or "").strip().upper() == "REGLER")
    en_cours = total - regles

    avec_tiers = sum(1 for r in records if "AVEC" in str(r.get("avec_sans_tiers") or "").upper()
                     and "SANS" not in str(r.get("avec_sans_tiers") or "").upper())
    sans_tiers = sum(1 for r in records if "SANS" in str(r.get("avec_sans_tiers") or "").upper())
    autre_tiers = total - (avec_tiers + sans_tiers)

    vehicules_conduits = sorted(list({str(r.get("immatriculation") or r.get("matricule_vehicule") or "").strip()
                                      for r in records if r.get("immatriculation") or r.get("matricule_vehicule")}))

    return {
        "chauffeur": chauffeur_name,
        "total": total,
        "regles": regles,
        "en_cours": en_cours,
        "avec_tiers": avec_tiers,
        "sans_tiers": sans_tiers,
        "autre_tiers": autre_tiers,
        "vehicules": [v for v in vehicules_conduits if v],
        "records": records,
    }


def get_vehicle_statistics(immatriculation):
    """Calcule les statistiques et récupère l'historique d'un véhicule donné."""
    norm_target = _normalize_name(immatriculation)
    all_records = db.fetch_all()
    records = [r for r in all_records
               if _normalize_name(r.get("immatriculation")) == norm_target
               or _normalize_name(r.get("matricule_vehicule")) == norm_target]

    def _sort_key(r):
        dt = date_utils.parse_date_input(r.get("date_sinistre"))
        return dt or datetime.date(1970, 1, 1)

    records.sort(key=_sort_key, reverse=True)
    total = len(records)
    regles = sum(1 for r in records if str(r.get("statut_reglement") or "").strip().upper() == "REGLER")
    chauffeurs = sorted(list({str(r.get("chauffeur") or "").strip() for r in records if r.get("chauffeur")}))

    return {
        "immatriculation": immatriculation,
        "total": total,
        "regles": regles,
        "en_cours": total - regles,
        "chauffeurs": [c for c in chauffeurs if c],
        "records": records,
    }


def export_releve_chauffeur_word(chauffeur_name, output_path=None):
    """Génère un rapport officiel au format Word (.docx) :
    « RELEVÉ INDIVIDUEL DE SINISTRALITÉ CHAUFFEUR ».

    Comprend :
    - En-tête officiel de l'organisme.
    - Synthèse statistique du chauffeur (total, réglés, implication tiers).
    - Tableau chronologique complet des sinistres.
    - Zone de signature de la Direction / Gestion de flotte.
    """
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    stats = get_driver_statistics(chauffeur_name)
    if output_path is None:
        import re
        safe_name = re.sub(r'[^a-zA-Z0-9_\-]', '_', chauffeur_name).strip('_') or "Chauffeur"
        default_dir = os.path.join(db.get_app_dir(), "Fiches")
        os.makedirs(default_dir, exist_ok=True)
        output_path = os.path.join(default_dir, f"Releve_Sinistralite_{safe_name}.docx")

    doc = docx.Document()

    # Marges de page
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # 1. En-tête Organisme
    letterhead = fiche.get_letterhead()
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, line in enumerate(letterhead):
        run = p_org.add_run(line + ("\n" if i < len(letterhead) - 1 else ""))
        run.font.name = "Arial"
        if i == 0:
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(31, 58, 95)
        else:
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    # 2. Titre principal
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("RELEVÉ INDIVIDUEL DE SINISTRALITÉ")
    r_title.bold = True
    r_title.font.name = "Arial"
    r_title.font.size = Pt(16)
    r_title.font.color.rgb = RGBColor(31, 58, 95)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run(f"Chauffeur / Conducteur : {chauffeur_name}")
    r_sub.bold = True
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(50, 50, 50)

    doc.add_paragraph()

    # 3. Tableau de synthèse KPI (2 colonnes)
    p_s = doc.add_paragraph()
    r_s = p_s.add_run("1. Synthèse statistique")
    r_s.bold = True
    r_s.font.name = "Arial"
    r_s.font.size = Pt(12)
    r_s.font.color.rgb = RGBColor(31, 58, 95)

    table_kpi = doc.add_table(rows=4, cols=2)
    table_kpi.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_kpi.style = "Table Grid"

    kpis_rows = [
        ("Nombre total de sinistres enregistrés", str(stats["total"])),
        ("Dossiers réglés / clôturés", f"{stats['regles']} ({stats['en_cours']} en cours/instance)"),
        ("Implication avec tiers reconnu", f"{stats['avec_tiers']} dossier(s) avec tiers"),
        ("Sinistres sans tiers (responsabilité unique)", f"{stats['sans_tiers']} dossier(s) sans tiers"),
    ]

    for i, (label, val) in enumerate(kpis_rows):
        row = table_kpi.rows[i]
        row.height = Pt(22)
        row.cells[0].width = Inches(4.2)
        row.cells[1].width = Inches(2.5)

        c0 = row.cells[0].paragraphs[0]
        r0 = c0.add_run(label)
        r0.font.name = "Arial"
        r0.font.size = Pt(10)

        c1 = row.cells[1].paragraphs[0]
        r1 = c1.add_run(val)
        r1.bold = True
        r1.font.name = "Arial"
        r1.font.size = Pt(10)

    doc.add_paragraph()

    # 4. Tableau chronologique des sinistres
    p_d = doc.add_paragraph()
    r_d = p_d.add_run("2. Historique détaillé des sinistres")
    r_d.bold = True
    r_d.font.name = "Arial"
    r_d.font.size = Pt(12)
    r_d.font.color.rgb = RGBColor(31, 58, 95)

    cols_headers = ["N° Dossier", "Date", "Véhicule (Matricule)", "Tiers", "Dégâts déclarés", "Statut"]
    col_widths = [Inches(1.0), Inches(0.9), Inches(1.2), Inches(1.0), Inches(1.8), Inches(0.8)]

    table_hist = doc.add_table(rows=1 + max(1, len(stats["records"])), cols=len(cols_headers))
    table_hist.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_hist.style = "Table Grid"

    # En-tête du tableau
    hdr_row = table_hist.rows[0]
    for idx, (th_str, w) in enumerate(zip(cols_headers, col_widths)):
        cell = hdr_row.cells[idx]
        cell.width = w
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(th_str)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(9.5)

    if not stats["records"]:
        # Aucune donnée
        row = table_hist.rows[1]
        row.cells[0].paragraphs[0].add_run("Aucun sinistre répertorié")
    else:
        for idx, rec in enumerate(stats["records"]):
            row = table_hist.rows[idx + 1]
            row.height = Pt(20)
            for col_idx, w in enumerate(col_widths):
                row.cells[col_idx].width = w

            num_dos = str(rec.get("numero_dossier") or rec.get("numero") or rec.get("id") or "-")
            dt_str = str(rec.get("date_sinistre") or "-")
            veh = str(rec.get("immatriculation") or rec.get("matricule_vehicule") or rec.get("type_vehicule") or "-")
            tiers = str(rec.get("avec_sans_tiers") or "-")
            deg = str(rec.get("degats_cause") or "-")
            if len(deg) > 40:
                deg = deg[:37] + "..."
            statut = str(rec.get("statut_reglement") or "-")

            row_data = [num_dos, dt_str, veh, tiers, deg, statut]
            for c_i, text_val in enumerate(row_data):
                p = row.cells[c_i].paragraphs[0]
                r = p.add_run(str(text_val))
                r.font.name = "Arial"
                r.font.size = Pt(8.5)

    doc.add_paragraph()
    doc.add_paragraph()

    # 5. Zone de signature
    table_sig = doc.add_table(rows=1, cols=2)
    table_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_sig.style = "Table Grid"
    table_sig.rows[0].height = Inches(1.1)

    now_str = datetime.datetime.now().strftime("%d/%m/%Y")
    sigs = [
        f"Établi le {now_str}\n\nLe Gestionnaire de Flotte",
        "Visa de la Direction / RH"
    ]
    for i, sig_txt in enumerate(sigs):
        cell = table_sig.rows[0].cells[i]
        cell.width = Inches(3.3)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(sig_txt)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(9.5)

    doc.save(output_path)
    return output_path
